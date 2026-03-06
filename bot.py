import asyncio
import hashlib
import json
import logging
import os
import re
import sqlite3
import shutil
import tempfile
import time
import uuid
import zipfile
from collections import defaultdict, deque
from pathlib import Path
from typing import Deque, Dict, List, Optional, Set, Tuple

import pypdfium2 as pdfium
import pytesseract
from pdf2docx import Converter
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from PIL import Image, ImageOps
from telegram.error import NetworkError, TimedOut
from telegram import LabeledPrice, Update
from telegram.constants import ChatAction
from telegram.ext import (
    Application,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    PreCheckoutQueryHandler,
    filters,
)

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}
SUPPORTED_MIME_PREFIX = ("image/", "application/pdf")
FREE_MAX_PAGES = 5
FREE_MAX_INPUT_BYTES = 10 * 1024 * 1024
DEFAULT_MAX_CONCURRENT_JOBS = 2
DEFAULT_OCR_TIMEOUT_SECONDS = 120
DEFAULT_TEXT_NATIVE_MIN_CHARS = 40
DEFAULT_CACHE_DIR = "cache"
DEFAULT_DB_PATH = "data/billing.sqlite3"
DEFAULT_PREMIUM_DAYS = 30
DEFAULT_PREMIUM_PRICE_USD_CENTS = 500
PREMIUM_INVOICE_PAYLOAD = "premium_monthly_subscription"
DEFAULT_FREE_REQUESTS_PER_MINUTE = 3
DEFAULT_PREMIUM_REQUESTS_PER_MINUTE = 20
DEFAULT_RATE_LIMIT_WINDOW_SECONDS = 60
DEFAULT_CACHE_TTL_DAYS = 7
DEFAULT_CACHE_MAX_SIZE_MB = 1024
DEFAULT_JOB_HISTORY_PER_USER = 10
DEFAULT_PDF_LAYOUT_MIN_NATIVE_RATIO = 0.8
DEFAULT_RUN_MODE = "polling"
DEFAULT_WEBHOOK_PORT = 8080
DEFAULT_WEBHOOK_PATH = "/telegram/webhook"
DEFAULT_WEBHOOK_LISTEN = "0.0.0.0"
DEFAULT_TELEGRAM_WRITE_TIMEOUT_SECONDS = 180
DEFAULT_MAX_OUTPUT_DOCX_MB = 25
DEFAULT_ENABLE_DOCX_MEDIA_COMPRESSION = True
DEFAULT_DOCX_MEDIA_COMPRESSION_MIN_MB = 8
DEFAULT_DOCX_IMAGE_MAX_DIMENSION = 1800
DEFAULT_DOCX_IMAGE_JPEG_QUALITY = 70
DEFAULT_AUTO_OCR_LANG_CANDIDATES = "vie+eng,eng"
DEFAULT_AUTO_OCR_DETECT_SAMPLE_PAGES = 2


def _log_event(level: int, event: str, **fields: object) -> None:
    payload = {"event": event, **fields}
    logger.log(level, json.dumps(payload, ensure_ascii=True, default=str))


def _configure_tesseract() -> None:
    tesseract_cmd = os.getenv("TESSERACT_CMD", "").strip()
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd


def _init_billing_db(db_path: Path) -> None:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                user_id INTEGER PRIMARY KEY,
                premium_until INTEGER,
                ocr_lang TEXT,
                updated_at INTEGER NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS payments (
                payment_id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                telegram_charge_id TEXT,
                provider_charge_id TEXT,
                currency TEXT,
                total_amount INTEGER,
                invoice_payload TEXT,
                paid_at INTEGER NOT NULL
            )
            """
        )

        # Backward-compatible migration for existing DB files.
        cols = {row[1] for row in conn.execute("PRAGMA table_info(users)")}
        if "ocr_lang" not in cols:
            conn.execute("ALTER TABLE users ADD COLUMN ocr_lang TEXT")
        conn.commit()


def _unix_ts() -> int:
    return int(time.time())


def _get_user_premium_until(db_path: Path, user_id: int) -> Optional[int]:
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT premium_until FROM users WHERE user_id = ?",
            (user_id,),
        ).fetchone()
    return int(row[0]) if row and row[0] is not None else None


def _get_user_ocr_lang(db_path: Path, user_id: int) -> Optional[str]:
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT ocr_lang FROM users WHERE user_id = ?",
            (user_id,),
        ).fetchone()
    if not row or row[0] is None:
        return None
    value = str(row[0]).strip()
    return value or None


def _set_user_ocr_lang(db_path: Path, user_id: int, ocr_lang: str) -> None:
    now = _unix_ts()
    existing_premium_until = _get_user_premium_until(db_path, user_id)
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO users (user_id, premium_until, ocr_lang, updated_at)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(user_id) DO UPDATE SET
                ocr_lang = excluded.ocr_lang,
                updated_at = excluded.updated_at
            """,
            (user_id, existing_premium_until, ocr_lang, now),
        )
        conn.commit()


def _record_payment(
    db_path: Path,
    user_id: int,
    telegram_charge_id: str,
    provider_charge_id: str,
    currency: str,
    total_amount: int,
    invoice_payload: str,
) -> None:
    now = _unix_ts()
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO payments (
                user_id,
                telegram_charge_id,
                provider_charge_id,
                currency,
                total_amount,
                invoice_payload,
                paid_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                user_id,
                telegram_charge_id,
                provider_charge_id,
                currency,
                total_amount,
                invoice_payload,
                now,
            ),
        )
        conn.commit()


def _is_user_premium(db_path: Path, user_id: int) -> bool:
    premium_until = _get_user_premium_until(db_path, user_id)
    return premium_until is not None and premium_until > _unix_ts()


def _grant_premium_days(db_path: Path, user_id: int, days: int) -> int:
    now = _unix_ts()
    base = _get_user_premium_until(db_path, user_id) or now
    if base < now:
        base = now
    premium_until = base + days * 24 * 60 * 60

    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO users (user_id, premium_until, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(user_id) DO UPDATE SET
                premium_until = excluded.premium_until,
                updated_at = excluded.updated_at
            """,
            (user_id, premium_until, now),
        )
        conn.commit()

    return premium_until


def _revoke_premium(db_path: Path, user_id: int) -> None:
    now = _unix_ts()
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO users (user_id, premium_until, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(user_id) DO UPDATE SET
                premium_until = excluded.premium_until,
                updated_at = excluded.updated_at
            """,
            (user_id, now - 1, now),
        )
        conn.commit()


def _get_billing_stats(db_path: Path) -> Tuple[int, int]:
    now = _unix_ts()
    with sqlite3.connect(db_path) as conn:
        total_users = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        premium_active = conn.execute(
            "SELECT COUNT(*) FROM users WHERE premium_until IS NOT NULL AND premium_until > ?",
            (now,),
        ).fetchone()[0]
    return int(total_users), int(premium_active)


def _format_utc_date(unix_ts: int) -> str:
    return time.strftime("%Y-%m-%d %H:%M UTC", time.gmtime(unix_ts))


def _get_user_limits(is_premium: bool, is_admin: bool = False) -> Tuple[Optional[int], Optional[int], str]:
    if is_admin:
        return None, None, "Admin"
    if is_premium:
        return None, None, "Premium"
    return FREE_MAX_PAGES, FREE_MAX_INPUT_BYTES, "Free"


def _parse_admin_user_ids(raw: str) -> Set[int]:
    values = set()
    for token in raw.split(","):
        item = token.strip()
        if not item:
            continue
        try:
            values.add(int(item))
        except ValueError:
            logger.warning("Ignoring invalid ADMIN_USER_IDS entry: %s", item)
    return values


def _cleanup_cache_dir(cache_dir: Path, ttl_days: int, max_size_bytes: int) -> Tuple[int, int, int]:
    if not cache_dir.exists():
        return 0, 0, 0

    now = time.time()
    ttl_seconds = max(1, ttl_days) * 24 * 60 * 60
    removed_expired = 0
    removed_size = 0

    entries: List[Tuple[float, Path, int]] = []
    for entry in cache_dir.glob("*.docx"):
        try:
            stat = entry.stat()
        except OSError:
            continue
        mtime = stat.st_mtime
        size = stat.st_size
        if (now - mtime) > ttl_seconds:
            try:
                entry.unlink()
                removed_expired += 1
            except OSError:
                pass
            continue
        entries.append((mtime, entry, size))

    total_size = sum(size for _, _, size in entries)
    if max_size_bytes > 0 and total_size > max_size_bytes:
        for _, file_path, size in sorted(entries, key=lambda x: x[0]):
            if total_size <= max_size_bytes:
                break
            try:
                file_path.unlink()
                total_size -= size
                removed_size += 1
            except OSError:
                continue

    return removed_expired, removed_size, total_size


def _increment_stat(application: Application, key: str) -> None:
    stats = application.bot_data.get("runtime_stats")
    if not isinstance(stats, dict):
        stats = {}
        application.bot_data["runtime_stats"] = stats
    stats[key] = int(stats.get(key, 0)) + 1


async def _check_and_apply_rate_limit(
    application: Application,
    user_id: int,
    is_premium: bool,
) -> Tuple[bool, int]:
    if user_id <= 0:
        return True, 0

    lock: asyncio.Lock = application.bot_data["rate_limit_lock"]
    buckets: Dict[int, Deque[float]] = application.bot_data["rate_limit_buckets"]
    window_seconds = int(
        application.bot_data.get("rate_limit_window_seconds", DEFAULT_RATE_LIMIT_WINDOW_SECONDS)
    )
    free_rpm = int(application.bot_data.get("free_requests_per_minute", DEFAULT_FREE_REQUESTS_PER_MINUTE))
    premium_rpm = int(
        application.bot_data.get("premium_requests_per_minute", DEFAULT_PREMIUM_REQUESTS_PER_MINUTE)
    )
    allowed_requests = premium_rpm if is_premium else free_rpm

    now = time.time()
    async with lock:
        queue = buckets.get(user_id)
        if queue is None:
            queue = deque()
            buckets[user_id] = queue

        while queue and (now - queue[0]) >= window_seconds:
            queue.popleft()

        if len(queue) >= allowed_requests:
            retry_after = int(max(1, window_seconds - (now - queue[0])))
            return False, retry_after

        queue.append(now)
        return True, 0


def _preprocess_image(image: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(image)
    enhanced = ImageOps.autocontrast(gray)
    return enhanced


def _extract_images_from_pdf(pdf_path: Path) -> List[Image.Image]:
    images: List[Image.Image] = []
    pdf = pdfium.PdfDocument(str(pdf_path))
    try:
        page_count = len(pdf)
        page_count = min(page_count, FREE_MAX_PAGES)
        for page_index in range(page_count):
            page = pdf[page_index]
            bitmap = page.render(scale=2.0)
            pil_image = bitmap.to_pil().convert("RGB")
            images.append(pil_image)
    finally:
        pdf.close()
    return images


def _extract_text_from_pdf(pdf_path: Path, max_pages: Optional[int]) -> List[Tuple[int, str]]:
    pages: List[Tuple[int, str]] = []
    reader = PdfReader(str(pdf_path))
    total = len(reader.pages)
    if max_pages is not None:
        total = min(total, max_pages)

    for i in range(total):
        raw_text = reader.pages[i].extract_text() or ""
        pages.append((i + 1, raw_text.strip()))

    return pages


def _get_pdf_page_count(pdf_path: Path) -> int:
    reader = PdfReader(str(pdf_path))
    return len(reader.pages)


def _extract_selected_images_from_pdf(
    pdf_path: Path,
    page_numbers: List[int],
    max_pages: Optional[int],
) -> List[Tuple[int, Image.Image]]:
    images: List[Tuple[int, Image.Image]] = []
    pdf = pdfium.PdfDocument(str(pdf_path))
    try:
        page_count = len(pdf)
        if max_pages is not None:
            page_count = min(page_count, max_pages)
        for page_number in sorted(set(page_numbers)):
            index = page_number - 1
            if index < 0 or index >= page_count:
                continue
            page = pdf[index]
            bitmap = page.render(scale=2.0)
            pil_image = bitmap.to_pil().convert("RGB")
            images.append((page_number, pil_image))
    finally:
        pdf.close()
    return images


def _convert_pdf_to_docx_layout(
    pdf_path: Path,
    output_path: Path,
    max_pages: Optional[int],
) -> None:
    converter = Converter(str(pdf_path))
    try:
        end_page = max_pages if max_pages is not None else None
        converter.convert(str(output_path), start=0, end=end_page)
    finally:
        converter.close()


def _has_meaningful_text(text: str, min_chars: int) -> bool:
    cleaned = (text or "").strip()
    if len(cleaned) < min_chars:
        return False

    alnum_count = sum(1 for ch in cleaned if ch.isalnum())
    ratio = alnum_count / len(cleaned) if cleaned else 0.0
    return ratio >= 0.25


def _ocr_images(images: List[Image.Image], lang: str) -> List[Tuple[int, str]]:
    page_texts: List[Tuple[int, str]] = []
    for idx, image in enumerate(images, start=1):
        prepared = _preprocess_image(image)
        text = pytesseract.image_to_string(prepared, lang=lang, config="--oem 3 --psm 3")
        page_texts.append((idx, text.strip()))
    return page_texts


def _ocr_page_images(page_images: List[Tuple[int, Image.Image]], lang: str) -> List[Tuple[int, str]]:
    page_texts: List[Tuple[int, str]] = []
    for page_number, image in page_images:
        prepared = _preprocess_image(image)
        text = pytesseract.image_to_string(prepared, lang=lang, config="--oem 3 --psm 3")
        page_texts.append((page_number, text.strip()))
    return page_texts


def _score_ocr_text_quality(text: str) -> int:
    if not text:
        return 0
    cleaned = text.strip()
    if not cleaned:
        return 0

    alnum_count = sum(1 for ch in cleaned if ch.isalnum())
    vietnamese_diacritics = sum(
        1
        for ch in cleaned
        if ch in "ăâđêôơưĂÂĐÊÔƠƯáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ"
    )
    # Favor readable OCR output and slightly prefer preserved Vietnamese diacritics.
    return alnum_count + (2 * vietnamese_diacritics)


def _parse_auto_lang_candidates(raw_value: str) -> List[str]:
    candidates: List[str] = []
    for item in (raw_value or "").split(","):
        value = item.strip()
        if not value or value.lower() == "auto":
            continue
        if _is_valid_ocr_lang(value):
            candidates.append(value)

    if not candidates:
        return ["vie+eng", "eng"]
    return candidates


def _detect_ocr_lang_from_page_images(
    page_images: List[Tuple[int, Image.Image]],
    fallback_lang: str,
    candidate_langs: List[str],
    sample_pages: int,
) -> str:
    if not page_images:
        return fallback_lang

    sampled = page_images[: max(1, sample_pages)]
    prepared_images = [_preprocess_image(img) for _, img in sampled]

    best_lang = fallback_lang
    best_score = -1

    for candidate in candidate_langs:
        total_score = 0
        success = False
        for prepared in prepared_images:
            try:
                text = pytesseract.image_to_string(prepared, lang=candidate, config="--oem 3 --psm 6")
                total_score += _score_ocr_text_quality(text)
                success = True
            except Exception:
                continue

        if success and total_score > best_score:
            best_score = total_score
            best_lang = candidate

    return best_lang


def _build_docx(page_texts: List[Tuple[int, str]], output_path: Path) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for page_number, text in page_texts:
        document.add_heading(f"Page {page_number}", level=2)
        if text:
            for paragraph in text.splitlines():
                if paragraph.strip():
                    document.add_paragraph(paragraph)
                else:
                    document.add_paragraph("")
        else:
            document.add_paragraph("[No text could be recognized on this page]")
        document.add_page_break()

    # Remove trailing empty page by deleting the final page break paragraph when possible.
    if document.paragraphs:
        last = document.paragraphs[-1]
        if not last.text:
            p = last._element
            p.getparent().remove(p)

    document.save(output_path)


def _compress_docx_media(
    docx_path: Path,
    max_dimension: int,
    jpeg_quality: int,
) -> Tuple[bool, int, int]:
    if not docx_path.exists():
        return False, 0, 0

    original_size = docx_path.stat().st_size
    changed_files = 0

    with tempfile.TemporaryDirectory(prefix="docx_media_compress_") as tmp:
        base_dir = Path(tmp)
        extract_dir = base_dir / "extract"
        extract_dir.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(extract_dir)

        media_dir = extract_dir / "word" / "media"
        if media_dir.exists():
            for media_file in media_dir.iterdir():
                if not media_file.is_file():
                    continue
                suffix = media_file.suffix.lower()
                if suffix not in {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}:
                    continue

                try:
                    with Image.open(media_file) as img:
                        image = ImageOps.exif_transpose(img)
                        if image.mode not in {"RGB", "L"}:
                            image = image.convert("RGB")

                        width, height = image.size
                        max_side = max(width, height)
                        if max_side > max_dimension:
                            scale = max_dimension / max_side
                            new_size = (max(1, int(width * scale)), max(1, int(height * scale)))
                            image = image.resize(new_size, Image.Resampling.LANCZOS)

                        if suffix in {".jpg", ".jpeg"}:
                            image.save(media_file, format="JPEG", quality=jpeg_quality, optimize=True)
                        elif suffix == ".png":
                            image.save(media_file, format="PNG", optimize=True)
                        else:
                            image.save(media_file)

                    changed_files += 1
                except Exception:
                    continue

        if changed_files == 0:
            return False, original_size, original_size

        compressed_docx = base_dir / "compressed.docx"
        with zipfile.ZipFile(compressed_docx, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for file in extract_dir.rglob("*"):
                if file.is_file():
                    arcname = file.relative_to(extract_dir).as_posix()
                    zout.write(file, arcname=arcname)

        new_size = compressed_docx.stat().st_size
        if new_size < original_size:
            shutil.copyfile(compressed_docx, docx_path)
            return True, original_size, new_size

        return False, original_size, new_size


def _is_supported_document(filename: str | None, mime_type: str | None) -> bool:
    if filename:
        suffix = Path(filename).suffix.lower()
        if suffix in SUPPORTED_EXTENSIONS:
            return True
    if mime_type and any(mime_type.startswith(prefix) for prefix in SUPPORTED_MIME_PREFIX):
        return True
    return False


def _sanitize_output_stem(raw_name: str) -> str:
    # Keep a conservative filename charset for cross-platform compatibility.
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", (raw_name or "").strip())
    cleaned = cleaned.strip("._-")
    return cleaned or "output"


def _safe_unlink(file_path: Path | None) -> None:
    if not file_path:
        return
    try:
        if file_path.exists():
            file_path.unlink()
    except Exception as exc:
        logger.warning("Cannot delete temp file %s: %s", file_path, exc)


def _close_images(images: List[Image.Image]) -> None:
    for img in images:
        try:
            img.close()
        except Exception:
            pass


def _close_page_images(page_images: List[Tuple[int, Image.Image]]) -> None:
    for _, img in page_images:
        try:
            img.close()
        except Exception:
            pass


async def _run_in_thread_with_timeout(func, *args, timeout_seconds: int):
    return await asyncio.wait_for(asyncio.to_thread(func, *args), timeout=timeout_seconds)


async def _reply_document_with_retry(
    message,
    file_path: Path,
    filename: str,
    caption: str,
    write_timeout: float,
    retries: int = 2,
) -> None:
    for attempt in range(1, retries + 1):
        try:
            with file_path.open("rb") as f:
                await message.reply_document(
                    document=f,
                    filename=filename,
                    caption=caption,
                    write_timeout=write_timeout,
                    read_timeout=write_timeout,
                    connect_timeout=30,
                )
            return
        except (TimedOut, NetworkError) as exc:
            if attempt >= retries:
                raise
            _log_event(
                logging.WARNING,
                "send_document_retry",
                attempt=attempt,
                retries=retries,
                filename=filename,
                error=str(exc),
            )
            await asyncio.sleep(2 * attempt)


def _compute_file_sha256(file_path: Path) -> str:
    digest = hashlib.sha256()
    with file_path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Send a scanned file (PDF/JPG/PNG/...) and I will OCR it to DOCX.\n"
        "Note: text-native PDFs are converted directly, scanned pages use OCR.\n"
        "Free plan limits: 5 pages and 10MB. Use /buy to subscribe Premium for $5/month (no page/file limits)."
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "How to use:\n"
        "1) Send an image or scanned PDF\n"
        "2) Wait while OCR and DOCX conversion run\n"
        "3) Receive the output .docx file\n"
        "Commands: /plan, /buy, /lang, /status"
    )


def _is_valid_ocr_lang(value: str) -> bool:
    if value.lower() == "auto":
        return True
    return re.fullmatch(r"[a-zA-Z+_\-]{2,30}", value) is not None


def _set_job_status(
    application: Application,
    user_id: int,
    job_id: str,
    status: str,
    detail: str,
) -> None:
    jobs_by_user: Dict[int, Dict[str, Dict[str, object]]] = application.bot_data["jobs_by_user"]
    job_order_by_user: Dict[int, Deque[str]] = application.bot_data["job_order_by_user"]
    per_user_history = int(application.bot_data.get("job_history_per_user", DEFAULT_JOB_HISTORY_PER_USER))
    now = _unix_ts()

    user_jobs = jobs_by_user.setdefault(user_id, {})
    if job_id not in user_jobs:
        order = job_order_by_user.setdefault(user_id, deque())
        order.append(job_id)
        while len(order) > per_user_history:
            evict_id = order.popleft()
            user_jobs.pop(evict_id, None)

    user_jobs[job_id] = {
        "status": status,
        "detail": detail,
        "updated_at": now,
    }


def _get_latest_job(application: Application, user_id: int) -> Optional[Tuple[str, Dict[str, object]]]:
    jobs_by_user: Dict[int, Dict[str, Dict[str, object]]] = application.bot_data["jobs_by_user"]
    job_order_by_user: Dict[int, Deque[str]] = application.bot_data["job_order_by_user"]
    order = job_order_by_user.get(user_id)
    if not order:
        return None
    job_id = order[-1]
    user_jobs = jobs_by_user.get(user_id, {})
    info = user_jobs.get(job_id)
    if not info:
        return None
    return job_id, info


async def lang_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    message = update.message
    if not user or not message:
        return

    db_path: Path = context.application.bot_data["billing_db_path"]
    if not context.args:
        current = _get_user_ocr_lang(db_path, user.id) or os.getenv("OCR_LANG", "vie+eng")
        await message.reply_text(
            "Usage: /lang <code>\n"
            "Example: /lang auto or /lang eng or /lang vie+eng\n"
            f"Current OCR language: {current}"
        )
        return

    lang_value = context.args[0].strip()
    if not _is_valid_ocr_lang(lang_value):
        await message.reply_text("Invalid language code format. Example: auto, eng, or vie+eng")
        return

    _set_user_ocr_lang(db_path, user.id, lang_value)
    await message.reply_text(f"OCR language updated to: {lang_value}")


async def status_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    message = update.message
    if not user or not message:
        return

    latest = _get_latest_job(context.application, user.id)
    if not latest:
        await message.reply_text("No recent job found. Send a file first.")
        return

    job_id, info = latest
    updated_at = int(info.get("updated_at", _unix_ts()))
    await message.reply_text(
        f"Job ID: {job_id}\n"
        f"Status: {info.get('status', 'unknown')}\n"
        f"Detail: {info.get('detail', '')}\n"
        f"Updated: {_format_utc_date(updated_at)}"
    )


def _is_admin_user(user_id: int, admin_user_ids: Set[int]) -> bool:
    return user_id in admin_user_ids


async def grant_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    message = update.message
    if not user or not message:
        return

    admin_user_ids: Set[int] = context.application.bot_data["admin_user_ids"]
    if not _is_admin_user(user.id, admin_user_ids):
        await message.reply_text("You are not authorized to use this command.")
        return

    if len(context.args) < 1:
        await message.reply_text("Usage: /grant <user_id> [days]")
        return

    try:
        target_user_id = int(context.args[0])
        days = int(context.args[1]) if len(context.args) > 1 else int(
            os.getenv("PREMIUM_DAYS", str(DEFAULT_PREMIUM_DAYS))
        )
        if days <= 0:
            raise ValueError("days must be positive")
    except ValueError:
        await message.reply_text("Invalid arguments. Usage: /grant <user_id> [days]")
        return

    db_path: Path = context.application.bot_data["billing_db_path"]
    premium_until = _grant_premium_days(db_path, target_user_id, days)
    await message.reply_text(
        f"Granted Premium to user {target_user_id} for {days} days.\n"
        f"Valid until: {_format_utc_date(premium_until)}"
    )


async def revoke_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    message = update.message
    if not user or not message:
        return

    admin_user_ids: Set[int] = context.application.bot_data["admin_user_ids"]
    if not _is_admin_user(user.id, admin_user_ids):
        await message.reply_text("You are not authorized to use this command.")
        return

    if len(context.args) != 1:
        await message.reply_text("Usage: /revoke <user_id>")
        return

    try:
        target_user_id = int(context.args[0])
    except ValueError:
        await message.reply_text("Invalid user_id. Usage: /revoke <user_id>")
        return

    db_path: Path = context.application.bot_data["billing_db_path"]
    _revoke_premium(db_path, target_user_id)
    await message.reply_text(f"Premium revoked for user {target_user_id}.")


async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    message = update.message
    if not user or not message:
        return

    admin_user_ids: Set[int] = context.application.bot_data["admin_user_ids"]
    if not _is_admin_user(user.id, admin_user_ids):
        await message.reply_text("You are not authorized to use this command.")
        return

    db_path: Path = context.application.bot_data["billing_db_path"]
    total_users, premium_active = _get_billing_stats(db_path)
    stats: Dict[str, int] = context.application.bot_data.get("runtime_stats", {})
    cache_dir: Path = context.application.bot_data["cache_dir"]
    cache_files = list(cache_dir.glob("*.docx"))
    cache_count = len(cache_files)
    cache_size_bytes = sum(f.stat().st_size for f in cache_files if f.exists())

    await message.reply_text(
        "Runtime stats\n"
        f"Total users in DB: {total_users}\n"
        f"Active Premium users: {premium_active}\n"
        f"Success: {int(stats.get('success', 0))}\n"
        f"Cache hits: {int(stats.get('cache_hit', 0))}\n"
        f"Rejected: {int(stats.get('rejected', 0))}\n"
        f"Timeout: {int(stats.get('timeout', 0))}\n"
        f"Failed: {int(stats.get('failed', 0))}\n"
        f"Cache files: {cache_count}\n"
        f"Cache size: {cache_size_bytes / (1024 * 1024):.2f} MB"
    )


async def plan_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    if not user:
        return

    db_path: Path = context.application.bot_data["billing_db_path"]
    admin_user_ids: Set[int] = context.application.bot_data["admin_user_ids"]
    if user.id in admin_user_ids:
        await update.message.reply_text(
            "Your plan: Admin\n"
            "Limits: No page/file-size/rate limits"
        )
        return

    premium_until = _get_user_premium_until(db_path, user.id)

    if premium_until and premium_until > _unix_ts():
        await update.message.reply_text(
            "Your plan: Premium\n"
            f"Valid until: {_format_utc_date(premium_until)}\n"
            "Limits: No page/file-size limits"
        )
    else:
        await update.message.reply_text(
            "Your plan: Free\n"
            "Limits: 5 pages, 10MB per file\n"
            "Use /buy to start a monthly subscription ($5/month)."
        )


async def buy_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    message = update.message
    if not user or not message:
        return

    provider_token = os.getenv("TELEGRAM_PROVIDER_TOKEN", "").strip()
    if not provider_token:
        await message.reply_text(
            "Payments are not configured yet. Set TELEGRAM_PROVIDER_TOKEN in .env to enable /buy."
        )
        return

    premium_days = int(os.getenv("PREMIUM_DAYS", str(DEFAULT_PREMIUM_DAYS)))
    premium_price = int(
        os.getenv("PREMIUM_PRICE_USD_CENTS", str(DEFAULT_PREMIUM_PRICE_USD_CENTS))
    )
    monthly_price_usd = premium_price / 100

    await context.bot.send_invoice(
        chat_id=message.chat_id,
        title="Scan2DOCX Premium Monthly",
        description=(
            f"Monthly subscription: ${monthly_price_usd:.2f}/month for unlimited file size and pages."
        ),
        payload=PREMIUM_INVOICE_PAYLOAD,
        provider_token=provider_token,
        currency="USD",
        prices=[LabeledPrice(label="Premium subscription (1 month)", amount=premium_price)],
        start_parameter="scan2docx-premium",
    )


async def precheckout_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.pre_checkout_query
    if not query:
        return

    if query.invoice_payload != PREMIUM_INVOICE_PAYLOAD:
        await query.answer(ok=False, error_message="Invalid payment payload.")
        return

    await query.answer(ok=True)


async def successful_payment_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.message
    user = update.effective_user
    if not message or not user or not message.successful_payment:
        return

    db_path: Path = context.application.bot_data["billing_db_path"]
    premium_days = int(os.getenv("PREMIUM_DAYS", str(DEFAULT_PREMIUM_DAYS)))
    premium_until = _grant_premium_days(db_path, user.id, premium_days)
    payment = message.successful_payment
    _record_payment(
        db_path=db_path,
        user_id=user.id,
        telegram_charge_id=payment.telegram_payment_charge_id,
        provider_charge_id=payment.provider_payment_charge_id,
        currency=payment.currency,
        total_amount=payment.total_amount,
        invoice_payload=payment.invoice_payload,
    )

    await message.reply_text(
        "Payment received. Your monthly Premium subscription is now active.\n"
        f"Valid until: {_format_utc_date(premium_until)}\n"
        "Limits removed: unlimited pages and file size. Renew monthly to keep Premium."
    )


async def _download_telegram_file(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    temp_dir: Path,
    max_input_bytes: Optional[int],
) -> Tuple[Path, str]:
    message = update.message

    if message.document:
        doc = message.document
        if not _is_supported_document(doc.file_name, doc.mime_type):
            raise ValueError("Unsupported file format.")
        if max_input_bytes is not None and doc.file_size and doc.file_size > max_input_bytes:
            max_mb = max_input_bytes // (1024 * 1024)
            raise ValueError(f"File exceeds the {max_mb}MB limit for your plan.")
        file_obj = await context.bot.get_file(doc.file_id)
        original_name = doc.file_name or "input.bin"
        src_suffix = Path(original_name).suffix.lower() or ".bin"
        src_path = temp_dir / f"source{src_suffix}"
        await file_obj.download_to_drive(custom_path=str(src_path))
        output_stem = _sanitize_output_stem(Path(original_name).stem)
        return src_path, output_stem

    if message.photo:
        photo = message.photo[-1]
        if max_input_bytes is not None and photo.file_size and photo.file_size > max_input_bytes:
            max_mb = max_input_bytes // (1024 * 1024)
            raise ValueError(f"Image exceeds the {max_mb}MB limit for your plan.")
        file_obj = await context.bot.get_file(photo.file_id)
        src_path = temp_dir / "source.jpg"
        await file_obj.download_to_drive(custom_path=str(src_path))
        output_stem = _sanitize_output_stem(f"photo_{photo.file_unique_id[:8]}")
        return src_path, output_stem

    raise ValueError("Only images and scanned PDF files are supported.")


async def process_scan(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.message
    if not message:
        return

    semaphore = context.application.bot_data.get("process_semaphore")
    if semaphore is None:
        semaphore = asyncio.Semaphore(DEFAULT_MAX_CONCURRENT_JOBS)
        context.application.bot_data["process_semaphore"] = semaphore

    queue_started_at = time.perf_counter()
    if semaphore.locked():
        await message.reply_text("The system is busy. Your request has been queued.")

    async with semaphore:
        request_id = uuid.uuid4().hex[:12]
        queue_wait_seconds = time.perf_counter() - queue_started_at
        await context.bot.send_chat_action(chat_id=message.chat_id, action=ChatAction.TYPING)

        user = update.effective_user
        db_path: Path = context.application.bot_data["billing_db_path"]
        user_id = user.id if user else 0
        default_ocr_lang = os.getenv("OCR_LANG", "vie+eng")
        ocr_lang = _get_user_ocr_lang(db_path, user_id) if user_id else None
        if not ocr_lang:
            ocr_lang = default_ocr_lang
        auto_lang_candidates = _parse_auto_lang_candidates(
            os.getenv("AUTO_OCR_LANG_CANDIDATES", DEFAULT_AUTO_OCR_LANG_CANDIDATES)
        )
        auto_lang_detect_sample_pages = int(
            os.getenv("AUTO_OCR_DETECT_SAMPLE_PAGES", str(DEFAULT_AUTO_OCR_DETECT_SAMPLE_PAGES))
        )
        ocr_timeout_seconds = int(os.getenv("OCR_TIMEOUT_SECONDS", str(DEFAULT_OCR_TIMEOUT_SECONDS)))
        text_native_min_chars = int(
            os.getenv("TEXT_NATIVE_MIN_CHARS", str(DEFAULT_TEXT_NATIVE_MIN_CHARS))
        )
        pdf_layout_min_native_ratio = float(
            os.getenv("PDF_LAYOUT_MIN_NATIVE_RATIO", str(DEFAULT_PDF_LAYOUT_MIN_NATIVE_RATIO))
        )
        telegram_write_timeout_seconds = float(
            os.getenv(
                "TELEGRAM_WRITE_TIMEOUT_SECONDS",
                str(DEFAULT_TELEGRAM_WRITE_TIMEOUT_SECONDS),
            )
        )
        enable_docx_media_compression = (
            os.getenv(
                "ENABLE_DOCX_MEDIA_COMPRESSION",
                "1" if DEFAULT_ENABLE_DOCX_MEDIA_COMPRESSION else "0",
            ).strip().lower()
            in {"1", "true", "yes", "on"}
        )
        docx_media_compression_min_bytes = int(
            float(
                os.getenv(
                    "DOCX_MEDIA_COMPRESSION_MIN_MB",
                    str(DEFAULT_DOCX_MEDIA_COMPRESSION_MIN_MB),
                )
            )
            * 1024
            * 1024
        )
        docx_image_max_dimension = int(
            os.getenv("DOCX_IMAGE_MAX_DIMENSION", str(DEFAULT_DOCX_IMAGE_MAX_DIMENSION))
        )
        docx_image_jpeg_quality = int(
            os.getenv("DOCX_IMAGE_JPEG_QUALITY", str(DEFAULT_DOCX_IMAGE_JPEG_QUALITY))
        )
        max_output_docx_bytes = int(
            float(os.getenv("MAX_OUTPUT_DOCX_MB", str(DEFAULT_MAX_OUTPUT_DOCX_MB))) * 1024 * 1024
        )
        if user_id:
            _set_job_status(context.application, user_id, request_id, "queued", "Request accepted")
        admin_user_ids: Set[int] = context.application.bot_data["admin_user_ids"]
        is_admin = user_id in admin_user_ids if user_id else False
        is_premium = _is_user_premium(db_path, user_id) if user_id else False
        if not is_admin:
            allowed, retry_after = await _check_and_apply_rate_limit(
                context.application,
                user_id,
                is_premium,
            )
            if not allowed:
                await message.reply_text(
                    f"Rate limit reached. Please try again in {retry_after} seconds."
                )
                _log_event(
                    logging.INFO,
                    "process_rejected_rate_limit",
                    request_id=request_id,
                    user_id=user_id,
                    retry_after_seconds=retry_after,
                )
                _increment_stat(context.application, "rejected")
                if user_id:
                    _set_job_status(context.application, user_id, request_id, "rejected", "Rate limit reached")
                return

        max_pages, max_input_bytes, plan_name = _get_user_limits(is_premium, is_admin=is_admin)
        _log_event(
            logging.INFO,
            "process_start",
            request_id=request_id,
            user_id=user_id,
            plan=plan_name,
            queue_wait_seconds=round(queue_wait_seconds, 3),
            max_pages=max_pages,
            max_input_bytes=max_input_bytes,
            ocr_lang=ocr_lang,
            auto_lang_candidates=auto_lang_candidates,
        )
        total_started_at = time.perf_counter()
        status = "failed"
        total_pages = 0
        native_pages = 0
        ocr_pages = 0
        ocr_time_seconds = 0.0
        source_hash = ""
        effective_ocr_lang = ocr_lang

        cache_dir = context.application.bot_data.get("cache_dir")
        if cache_dir is None:
            cache_dir = Path(DEFAULT_CACHE_DIR)
            cache_dir.mkdir(parents=True, exist_ok=True)
            context.application.bot_data["cache_dir"] = cache_dir

        now_ts = time.time()
        last_cleanup_ts = float(context.application.bot_data.get("last_cache_cleanup_ts", 0.0))
        if now_ts - last_cleanup_ts >= 300:
            ttl_days = int(context.application.bot_data.get("cache_ttl_days", DEFAULT_CACHE_TTL_DAYS))
            max_size_bytes = int(
                context.application.bot_data.get(
                    "cache_max_size_bytes",
                    DEFAULT_CACHE_MAX_SIZE_MB * 1024 * 1024,
                )
            )
            removed_expired, removed_size, kept_size = _cleanup_cache_dir(
                cache_dir,
                ttl_days,
                max_size_bytes,
            )
            context.application.bot_data["last_cache_cleanup_ts"] = now_ts
            logger.info(
                "cache_cleanup periodic removed_expired=%s removed_size=%s kept_size_mb=%.2f",
                removed_expired,
                removed_size,
                kept_size / (1024 * 1024),
            )

        with tempfile.TemporaryDirectory(prefix="scan2docx_") as tmp:
            temp_dir = Path(tmp)
            src_path: Path | None = None
            output_path: Path | None = None
            output_stem = "output"
            images: List[Image.Image] = []
            selected_images: List[Tuple[int, Image.Image]] = []
            layout_conversion_used = False
            fallback_page_texts: List[Tuple[int, str]] = []

            try:
                src_path, output_stem = await _download_telegram_file(
                    update,
                    context,
                    temp_dir,
                    max_input_bytes,
                )
                _log_event(
                    logging.INFO,
                    "input_downloaded",
                    request_id=request_id,
                    user_id=user_id,
                    source_suffix=src_path.suffix.lower(),
                    output_stem=output_stem,
                )
            except ValueError as exc:
                await message.reply_text(str(exc))
                _log_event(
                    logging.INFO,
                    "process_rejected_input",
                    request_id=request_id,
                    user_id=user_id,
                    reason=str(exc),
                )
                if user_id:
                    _set_job_status(context.application, user_id, request_id, "rejected", str(exc))
                return

            await context.bot.send_chat_action(chat_id=message.chat_id, action=ChatAction.UPLOAD_DOCUMENT)

            try:
                if user_id:
                    _set_job_status(context.application, user_id, request_id, "processing", "Converting file")

                suffix = src_path.suffix.lower()
                source_hash = await asyncio.to_thread(_compute_file_sha256, src_path)
                cache_path = cache_dir / f"{source_hash}.docx"
                if cache_path.exists():
                    cache_size = cache_path.stat().st_size
                    if enable_docx_media_compression and cache_size >= docx_media_compression_min_bytes:
                        compressed, old_size, new_size = await asyncio.to_thread(
                            _compress_docx_media,
                            cache_path,
                            docx_image_max_dimension,
                            docx_image_jpeg_quality,
                        )
                        if compressed:
                            cache_size = new_size
                            _log_event(
                                logging.INFO,
                                "cache_docx_media_compressed",
                                request_id=request_id,
                                user_id=user_id,
                                old_size_bytes=old_size,
                                new_size_bytes=new_size,
                            )

                    if cache_size > max_output_docx_bytes:
                        _log_event(
                            logging.INFO,
                            "cache_bypass_large",
                            request_id=request_id,
                            user_id=user_id,
                            cache_size_bytes=cache_size,
                            max_output_docx_bytes=max_output_docx_bytes,
                        )
                    else:
                        out_name = f"{output_stem}.docx"
                        await _reply_document_with_retry(
                            message,
                            cache_path,
                            out_name,
                            "",
                            telegram_write_timeout_seconds,
                        )
                        status = "cache_hit"
                        _log_event(
                            logging.INFO,
                            "cache_hit",
                            request_id=request_id,
                            user_id=user_id,
                            output_name=out_name,
                        )
                        if user_id:
                            _set_job_status(context.application, user_id, request_id, "done", "Completed")
                        return

                if suffix == ".pdf":
                    pdf_page_count = await asyncio.to_thread(_get_pdf_page_count, src_path)
                    if max_pages is not None and pdf_page_count > max_pages:
                        raise ValueError(
                            f"PDF has {pdf_page_count} pages, exceeding your plan limit of {max_pages}."
                        )

                    text_pages = await asyncio.to_thread(_extract_text_from_pdf, src_path, max_pages)
                    fallback_page_texts = text_pages
                    total_pages = len(text_pages)
                    missing_text_pages = [
                        page for page, text in text_pages if not _has_meaningful_text(text, text_native_min_chars)
                    ]
                    native_pages = total_pages - len(missing_text_pages)
                    native_ratio = (native_pages / total_pages) if total_pages else 0.0
                    _log_event(
                        logging.INFO,
                        "pdf_analyzed",
                        request_id=request_id,
                        user_id=user_id,
                        total_pages=total_pages,
                        native_pages=native_pages,
                        missing_text_pages=len(missing_text_pages),
                        native_ratio=round(native_ratio, 3),
                        layout_ratio_threshold=pdf_layout_min_native_ratio,
                    )

                    out_name = f"{output_stem}.docx"
                    output_path = temp_dir / out_name
                    converted_with_layout = False

                    # Prefer layout-preserving conversion when PDF is mostly text-native.
                    if total_pages > 0 and native_ratio >= pdf_layout_min_native_ratio:
                        try:
                            await asyncio.to_thread(
                                _convert_pdf_to_docx_layout,
                                src_path,
                                output_path,
                                max_pages,
                            )
                            converted_with_layout = True
                            layout_conversion_used = True
                            ocr_pages = 0
                            _log_event(
                                logging.INFO,
                                "pdf_layout_conversion_success",
                                request_id=request_id,
                                user_id=user_id,
                                output_name=out_name,
                            )
                        except Exception as layout_exc:
                            logger.warning("Layout conversion failed, fallback to OCR pipeline: %s", layout_exc)
                            _log_event(
                                logging.WARNING,
                                "pdf_layout_conversion_failed",
                                request_id=request_id,
                                user_id=user_id,
                                error=str(layout_exc),
                            )

                    if not converted_with_layout and missing_text_pages:
                        selected_images = await asyncio.to_thread(
                            _extract_selected_images_from_pdf,
                            src_path,
                            missing_text_pages,
                            max_pages,
                        )
                        if ocr_lang.lower() == "auto":
                            effective_ocr_lang = await asyncio.to_thread(
                                _detect_ocr_lang_from_page_images,
                                selected_images,
                                default_ocr_lang,
                                auto_lang_candidates,
                                auto_lang_detect_sample_pages,
                            )
                            _log_event(
                                logging.INFO,
                                "ocr_lang_auto_detected",
                                request_id=request_id,
                                user_id=user_id,
                                detected_lang=effective_ocr_lang,
                                candidates=auto_lang_candidates,
                                pages_sampled=min(len(selected_images), auto_lang_detect_sample_pages),
                            )
                        else:
                            effective_ocr_lang = ocr_lang

                        ocr_started_at = time.perf_counter()
                        ocr_result = await _run_in_thread_with_timeout(
                            _ocr_page_images,
                            selected_images,
                            effective_ocr_lang,
                            timeout_seconds=ocr_timeout_seconds,
                        )
                        ocr_time_seconds = time.perf_counter() - ocr_started_at
                        ocr_map = {page: text for page, text in ocr_result}
                        page_texts = [(page, text or ocr_map.get(page, "")) for page, text in text_pages]
                        ocr_pages = len(missing_text_pages)
                        _log_event(
                            logging.INFO,
                            "pdf_ocr_fallback_done",
                            request_id=request_id,
                            user_id=user_id,
                            ocr_pages=ocr_pages,
                            ocr_time_seconds=round(ocr_time_seconds, 3),
                        )
                    elif not converted_with_layout:
                        page_texts = text_pages
                        _log_event(
                            logging.INFO,
                            "pdf_text_native_build",
                            request_id=request_id,
                            user_id=user_id,
                            pages=total_pages,
                        )

                    if not converted_with_layout:
                        await asyncio.to_thread(_build_docx, page_texts, output_path)
                else:
                    source_image = Image.open(src_path)
                    image = source_image.convert("RGB")
                    source_image.close()
                    images = [image]

                    image_pages = [(1, image)]
                    if ocr_lang.lower() == "auto":
                        effective_ocr_lang = await asyncio.to_thread(
                            _detect_ocr_lang_from_page_images,
                            image_pages,
                            default_ocr_lang,
                            auto_lang_candidates,
                            auto_lang_detect_sample_pages,
                        )
                        _log_event(
                            logging.INFO,
                            "ocr_lang_auto_detected",
                            request_id=request_id,
                            user_id=user_id,
                            detected_lang=effective_ocr_lang,
                            candidates=auto_lang_candidates,
                            pages_sampled=1,
                        )
                    else:
                        effective_ocr_lang = ocr_lang

                    ocr_started_at = time.perf_counter()
                    page_texts = await _run_in_thread_with_timeout(
                        _ocr_images,
                        images,
                        effective_ocr_lang,
                        timeout_seconds=ocr_timeout_seconds,
                    )
                    ocr_time_seconds = time.perf_counter() - ocr_started_at
                    total_pages = 1
                    native_pages = 0
                    ocr_pages = 1
                    out_name = f"{output_stem}.docx"
                    output_path = temp_dir / out_name
                    await asyncio.to_thread(_build_docx, page_texts, output_path)

                if output_path and output_path.exists():
                    output_size = output_path.stat().st_size

                    if enable_docx_media_compression and output_size >= docx_media_compression_min_bytes:
                        compressed, old_size, new_size = await asyncio.to_thread(
                            _compress_docx_media,
                            output_path,
                            docx_image_max_dimension,
                            docx_image_jpeg_quality,
                        )
                        if compressed:
                            output_size = new_size
                            _log_event(
                                logging.INFO,
                                "docx_media_compressed",
                                request_id=request_id,
                                user_id=user_id,
                                old_size_bytes=old_size,
                                new_size_bytes=new_size,
                            )

                    if layout_conversion_used and output_size > max_output_docx_bytes:
                        _log_event(
                            logging.WARNING,
                            "layout_output_too_large",
                            request_id=request_id,
                            user_id=user_id,
                            output_size_bytes=output_size,
                            max_output_docx_bytes=max_output_docx_bytes,
                        )
                        await asyncio.to_thread(_build_docx, fallback_page_texts, output_path)
                        layout_conversion_used = False
                        output_size = output_path.stat().st_size if output_path.exists() else output_size

                output_size = output_path.stat().st_size if output_path and output_path.exists() else 0

                await asyncio.to_thread(shutil.copyfile, output_path, cache_path)

                caption_lang = effective_ocr_lang if ocr_pages > 0 else "n/a"
                caption = (
                    f"Done ({plan_name}). Total pages: {total_pages}. "
                    f"Native: {native_pages}. OCR: {ocr_pages}. Lang: {caption_lang}."
                )
                if output_size > max_output_docx_bytes:
                    caption += " Output is large; Telegram upload may be unstable on weak networks."

                await _reply_document_with_retry(
                    message,
                    output_path,
                    out_name,
                    caption,
                    telegram_write_timeout_seconds,
                )
                status = "success"
                _log_event(
                    logging.INFO,
                    "document_sent",
                    request_id=request_id,
                    user_id=user_id,
                    output_name=out_name,
                    output_size_bytes=output_size,
                )
                if user_id:
                    _set_job_status(context.application, user_id, request_id, "done", caption)

            except ValueError as exc:
                status = "rejected"
                await message.reply_text(str(exc))
                _log_event(
                    logging.INFO,
                    "process_rejected",
                    request_id=request_id,
                    user_id=user_id,
                    reason=str(exc),
                )
                if user_id:
                    _set_job_status(context.application, user_id, request_id, "rejected", str(exc))
            except asyncio.TimeoutError:
                status = "timeout"
                await message.reply_text(
                    "OCR timed out. Please send a smaller or clearer file."
                )
                _log_event(
                    logging.WARNING,
                    "process_timeout_ocr",
                    request_id=request_id,
                    user_id=user_id,
                )
                if user_id:
                    _set_job_status(context.application, user_id, request_id, "timeout", "OCR timed out")
            except TimedOut:
                status = "timeout"
                await message.reply_text(
                    "Upload timed out while sending DOCX to Telegram. Please retry or send a smaller file."
                )
                _log_event(
                    logging.WARNING,
                    "process_timeout_upload",
                    request_id=request_id,
                    user_id=user_id,
                )
                if user_id:
                    _set_job_status(
                        context.application,
                        user_id,
                        request_id,
                        "timeout",
                        "Upload timed out while sending DOCX",
                    )
            except NetworkError:
                status = "timeout"
                await message.reply_text(
                    "Network error while sending DOCX to Telegram. Please retry in a moment."
                )
                _log_event(
                    logging.WARNING,
                    "process_timeout_network",
                    request_id=request_id,
                    user_id=user_id,
                )
                if user_id:
                    _set_job_status(
                        context.application,
                        user_id,
                        request_id,
                        "timeout",
                        "Network error while sending DOCX",
                    )
            except Exception as exc:
                logger.exception("Failed to process OCR: %s", exc)
                await message.reply_text(
                    "An error occurred during OCR/input processing. Please check your file and Tesseract setup."
                )
                _log_event(
                    logging.ERROR,
                    "process_exception",
                    request_id=request_id,
                    user_id=user_id,
                    error=str(exc),
                )
                if user_id:
                    _set_job_status(context.application, user_id, request_id, "failed", str(exc))
            finally:
                _close_images(images)
                _close_page_images(selected_images)
                _safe_unlink(output_path)
                _safe_unlink(src_path)

                total_seconds = time.perf_counter() - total_started_at
                _log_event(
                    logging.INFO,
                    "process_done",
                    request_id=request_id,
                    status=status,
                    plan=plan_name,
                    user_id=user_id,
                    pages_total=total_pages,
                    pages_native=native_pages,
                    pages_ocr=ocr_pages,
                    queue_wait_seconds=round(queue_wait_seconds, 3),
                    ocr_time_seconds=round(ocr_time_seconds, 3),
                    total_time_seconds=round(total_seconds, 3),
                )
                _increment_stat(context.application, status)


async def on_error(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.exception("Unhandled error: %s", context.error)


def main() -> None:
    load_dotenv()
    _configure_tesseract()

    token = os.getenv("BOT_TOKEN", "").strip()
    if not token:
        raise RuntimeError("BOT_TOKEN is not configured in .env")

    application = Application.builder().token(token).build()
    max_concurrent_jobs = int(os.getenv("MAX_CONCURRENT_JOBS", str(DEFAULT_MAX_CONCURRENT_JOBS)))
    cache_dir = Path(os.getenv("CACHE_DIR", DEFAULT_CACHE_DIR)).resolve()
    billing_db_path = Path(os.getenv("BILLING_DB_PATH", DEFAULT_DB_PATH)).resolve()
    admin_user_ids = _parse_admin_user_ids(os.getenv("ADMIN_USER_IDS", ""))
    free_requests_per_minute = int(
        os.getenv("FREE_REQUESTS_PER_MINUTE", str(DEFAULT_FREE_REQUESTS_PER_MINUTE))
    )
    premium_requests_per_minute = int(
        os.getenv("PREMIUM_REQUESTS_PER_MINUTE", str(DEFAULT_PREMIUM_REQUESTS_PER_MINUTE))
    )
    rate_limit_window_seconds = int(
        os.getenv("RATE_LIMIT_WINDOW_SECONDS", str(DEFAULT_RATE_LIMIT_WINDOW_SECONDS))
    )
    cache_ttl_days = int(os.getenv("CACHE_TTL_DAYS", str(DEFAULT_CACHE_TTL_DAYS)))
    cache_max_size_mb = int(os.getenv("CACHE_MAX_SIZE_MB", str(DEFAULT_CACHE_MAX_SIZE_MB)))
    cache_max_size_bytes = cache_max_size_mb * 1024 * 1024
    job_history_per_user = int(os.getenv("JOB_HISTORY_PER_USER", str(DEFAULT_JOB_HISTORY_PER_USER)))
    run_mode = os.getenv("RUN_MODE", DEFAULT_RUN_MODE).strip().lower()
    webhook_port = int(os.getenv("WEBHOOK_PORT", str(DEFAULT_WEBHOOK_PORT)))
    webhook_listen = os.getenv("WEBHOOK_LISTEN", DEFAULT_WEBHOOK_LISTEN).strip() or DEFAULT_WEBHOOK_LISTEN
    webhook_path = os.getenv("WEBHOOK_PATH", DEFAULT_WEBHOOK_PATH).strip() or DEFAULT_WEBHOOK_PATH
    webhook_url_base = os.getenv("WEBHOOK_URL", "").strip()
    webhook_secret_token = os.getenv("WEBHOOK_SECRET_TOKEN", "").strip()

    cache_dir.mkdir(parents=True, exist_ok=True)
    _init_billing_db(billing_db_path)
    removed_expired, removed_size, kept_size = _cleanup_cache_dir(
        cache_dir,
        cache_ttl_days,
        cache_max_size_bytes,
    )
    logger.info(
        "cache_cleanup startup removed_expired=%s removed_size=%s kept_size_mb=%.2f",
        removed_expired,
        removed_size,
        kept_size / (1024 * 1024),
    )

    application.bot_data["process_semaphore"] = asyncio.Semaphore(max_concurrent_jobs)
    application.bot_data["cache_dir"] = cache_dir
    application.bot_data["billing_db_path"] = billing_db_path
    application.bot_data["admin_user_ids"] = admin_user_ids
    application.bot_data["rate_limit_buckets"] = defaultdict(deque)
    application.bot_data["rate_limit_lock"] = asyncio.Lock()
    application.bot_data["free_requests_per_minute"] = free_requests_per_minute
    application.bot_data["premium_requests_per_minute"] = premium_requests_per_minute
    application.bot_data["rate_limit_window_seconds"] = rate_limit_window_seconds
    application.bot_data["runtime_stats"] = {
        "success": 0,
        "cache_hit": 0,
        "rejected": 0,
        "timeout": 0,
        "failed": 0,
    }
    application.bot_data["cache_ttl_days"] = cache_ttl_days
    application.bot_data["cache_max_size_bytes"] = cache_max_size_bytes
    application.bot_data["last_cache_cleanup_ts"] = 0.0
    application.bot_data["jobs_by_user"] = {}
    application.bot_data["job_order_by_user"] = defaultdict(deque)
    application.bot_data["job_history_per_user"] = job_history_per_user

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("plan", plan_command))
    application.add_handler(CommandHandler("buy", buy_command))
    application.add_handler(CommandHandler("lang", lang_command))
    application.add_handler(CommandHandler("status", status_command))
    application.add_handler(CommandHandler("grant", grant_command))
    application.add_handler(CommandHandler("revoke", revoke_command))
    application.add_handler(CommandHandler("stats", stats_command))
    application.add_handler(PreCheckoutQueryHandler(precheckout_callback))
    application.add_handler(MessageHandler(filters.SUCCESSFUL_PAYMENT, successful_payment_callback))

    scan_filter = (
        filters.PHOTO
        | filters.Document.PDF
        | filters.Document.IMAGE
        | filters.Document.FileExtension("png")
        | filters.Document.FileExtension("jpg")
        | filters.Document.FileExtension("jpeg")
        | filters.Document.FileExtension("tif")
        | filters.Document.FileExtension("tiff")
        | filters.Document.FileExtension("bmp")
        | filters.Document.FileExtension("webp")
    )
    application.add_handler(MessageHandler(scan_filter, process_scan))

    application.add_error_handler(on_error)

    if run_mode == "webhook":
        if not webhook_url_base:
            raise RuntimeError("WEBHOOK_URL must be set when RUN_MODE=webhook")

        clean_path = webhook_path.lstrip("/")
        full_webhook_url = f"{webhook_url_base.rstrip('/')}/{clean_path}"
        logger.info(
            "Bot is running in webhook mode on %s:%s path=/%s",
            webhook_listen,
            webhook_port,
            clean_path,
        )
        application.run_webhook(
            listen=webhook_listen,
            port=webhook_port,
            url_path=clean_path,
            webhook_url=full_webhook_url,
            drop_pending_updates=True,
            secret_token=webhook_secret_token or None,
        )
    else:
        logger.info("Bot is running in polling mode")
        application.run_polling(drop_pending_updates=True)


if __name__ == "__main__":
    main()
